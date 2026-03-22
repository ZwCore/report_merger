import os
import re
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def get_target_heading_level(paragraphs, current_index):
    """
    Scan backwards from current_index to find the nearest Heading.
    Return that level + 1. If none found, return 1.
    """
    for i in range(current_index - 1, -1, -1):
        p = paragraphs[i]
        if '{{' in p.text and '}}' in p.text:
            continue
            
        style = p.style
        name = style.name
        style_id = style.style_id if style.style_id else ""
        
        lvl = None
        if name.startswith('Heading') or name.startswith('标题'):
            parts = name.split()
            if len(parts) >= 2 and parts[-1].isdigit():
                lvl = int(parts[-1])
        
        if lvl is None and style_id.startswith('Heading'):
            suffix = style_id[7:]
            if suffix.isdigit():
                lvl = int(suffix)
                
        if lvl is not None:
            return lvl + 1
    return 1

def merge_reports(template_path, input_folder, output_path):
    doc = Document(template_path)
    
    matches = []
    paragraphs = doc.paragraphs 
    
    for i, paragraph in enumerate(paragraphs):
        if '{{' in paragraph.text and '}}' in paragraph.text:
             regex = r"\{\{(.*?)\}\}"
             found = re.findall(regex, paragraph.text)
             level = get_target_heading_level(paragraphs, i)
             for key in found:
                 matches.append((paragraph, key, level))
    
    chunk_counter = 1
    
    for paragraph, key, target_level in matches:
        filename = f"{key}.docx"
        file_path = os.path.join(input_folder, filename)
        
        if os.path.exists(file_path):
            try:
                # Prepare altChunk
                # 1. Read blob
                with open(file_path, 'rb') as f:
                    blob = f.read()
                
                # 2. Add Part to Package
                # We need a unique URI for each chunk
                part_uri = PackURI(f"/word/afchunk{chunk_counter}.docx")
                chunk_counter += 1
                
                # Content Type for generic docx chunk
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
                # Note: Correct content type for altChunk pointing to valid docx is typically:
                # application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml
                # or just .document
                
                # Let's create the Part
                new_part = Part(part_uri, content_type, blob)
                
                # 3. Create Relationship
                # RT.AFIP usually not in simple constant list?
                # "http://schemas.openxmlformats.org/officeDocument/2006/relationships/aFChunk"
                rel_id = doc.part.relate_to(new_part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/aFChunk")
                
                # 4. Create w:altChunk element
                alt_chunk = OxmlElement('w:altChunk')
                alt_chunk.set(qn('r:id'), rel_id)
                
                # Insertion Logic
                if paragraph._element.getparent() is None:
                    continue

                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                
                # Insert Title
                title_p = doc.add_paragraph(key, style=f'Heading {target_level}')
                parent.insert(index, title_p._element)
                index += 1
                
                # Insert Chunk
                # altChunk is a block level element, friend of p
                parent.insert(index, alt_chunk)
                index += 1
                
                # Cleanup placeholder
                p = paragraph._element
                if p.getparent():
                    p.getparent().remove(p)
                    p._p = p._element = None

            except Exception as e:
                print(f"Error processing {filename}: {e}")
        else:
            if paragraph._element.getparent():
                paragraph.clear()
                run = paragraph.add_run("待补充")
                run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.save(output_path)
    return output_path
